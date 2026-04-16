"""
ИИ-помощник для автоматизации документооборота отдела улучшений
MVP Backend — FastAPI + Ollama + faster-whisper + python-docx
"""

import os, json, uuid, re
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# ── опциональные зависимости ──────────────────────────────────────────────────
try:
    import httpx
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False

try:
    from faster_whisper import WhisperModel
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── конфиг ────────────────────────────────────────────────────────────────────
OLLAMA_URL   = os.getenv("OLLAMA_URL",   "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:7b")       # или saiga, llama3.1
WHISPER_MODEL_SIZE = os.getenv("WHISPER_MODEL", "base")       # tiny/base/small
UPLOAD_DIR   = Path("uploads");  UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR   = Path("outputs");  OUTPUT_DIR.mkdir(exist_ok=True)

app = FastAPI(title="ИИ-помощник Lean", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_methods=["*"], allow_headers=["*"])

# ── Whisper (lazy init) ───────────────────────────────────────────────────────
_whisper: Optional[object] = None

def get_whisper():
    global _whisper
    if _whisper is None and WHISPER_AVAILABLE:
        _whisper = WhisperModel(WHISPER_MODEL_SIZE, device="cpu", compute_type="int8")
    return _whisper

# ── Ollama helper ─────────────────────────────────────────────────────────────
async def llm_generate(prompt: str, system: str = "") -> str:
    """Вызов Ollama API."""
    if not OLLAMA_AVAILABLE:
        raise HTTPException(503, "httpx не установлен")
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    async with httpx.AsyncClient(timeout=120) as client:
        r = await client.post(f"{OLLAMA_URL}/api/chat", json={
            "model": OLLAMA_MODEL,
            "messages": messages,
            "stream": False,
            "options": {"temperature": 0.3, "num_predict": 2000}
        })
        r.raise_for_status()
        return r.json()["message"]["content"]

# ── промпты ───────────────────────────────────────────────────────────────────
SYSTEM_PROTOCOL = """Ты — корпоративный ИИ-помощник специалиста по бережливому производству.
Твоя задача — извлекать структурированную информацию из текста сессий решения проблем и формировать протоколы.
Отвечай строго в формате JSON, без пояснений и markdown-блоков. Используй русский язык."""

SYSTEM_PASSPORT = """Ты — корпоративный ИИ-помощник специалиста по управлению проектами улучшений.
Твоя задача — формировать паспорта проектов по методологии Lean на основе описания инициативы.
Отвечай строго в формате JSON, без пояснений и markdown-блоков. Используй русский язык."""

PROMPT_PROTOCOL = """Проанализируй следующий текст сессии решения проблем и извлеки структурированные данные.

ТЕКСТ СЕССИИ:
{text}

Верни JSON строго такой структуры (без лишних полей):
{{
  "date": "дата сессии или 'не указана'",
  "participants": ["список участников"],
  "problem": "чёткое описание проблемы",
  "root_causes": ["список корневых причин"],
  "solutions": ["список предложенных решений"],
  "actions": [
    {{"task": "описание задачи", "responsible": "ответственный", "deadline": "срок"}}
  ],
  "expected_effect": "ожидаемый эффект от внедрения",
  "next_steps": "следующие шаги"
}}"""

PROMPT_PASSPORT = """На основе следующего описания инициативы сформируй паспорт проекта улучшений.

ОПИСАНИЕ ИНИЦИАТИВЫ:
{text}

Верни JSON строго такой структуры:
{{
  "project_name": "краткое наименование проекта",
  "full_name": "полное наименование",
  "category": "категория (Lean/6С/SMED/TPM/СОП/другое)",
  "goal": "цель проекта",
  "problem_description": "описание текущей проблемы",
  "current_state": "описание текущего состояния (AS-IS)",
  "target_state": "описание целевого состояния (TO-BE)",
  "economic_effect": "ожидаемый экономический эффект",
  "resources": "необходимые ресурсы",
  "timeline": "базовый срок реализации",
  "project_manager": "руководитель проекта (если указан)",
  "risks": ["основные риски"],
  "success_metrics": ["метрики успеха"]
}}"""

# ── модели запросов ───────────────────────────────────────────────────────────
class TextRequest(BaseModel):
    text: str
    doc_type: str = "protocol"   # protocol | passport

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: list[ChatMessage]
    doc_type: str = "protocol"

# ── endpoints ─────────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {
        "status": "ok",
        "ollama": OLLAMA_AVAILABLE,
        "whisper": WHISPER_AVAILABLE,
        "docx": DOCX_AVAILABLE,
        "model": OLLAMA_MODEL,
        "whisper_model": WHISPER_MODEL_SIZE
    }

@app.get("/models")
async def list_models():
    """Список доступных моделей Ollama."""
    if not OLLAMA_AVAILABLE:
        return {"models": []}
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get(f"{OLLAMA_URL}/api/tags")
            models = [m["name"] for m in r.json().get("models", [])]
            return {"models": models}
    except Exception:
        return {"models": []}

@app.post("/transcribe")
async def transcribe(file: UploadFile = File(...)):
    """Транскрибация аудиофайла через faster-whisper."""
    if not WHISPER_AVAILABLE:
        raise HTTPException(503,
            "faster-whisper не установлен. Установите: pip install faster-whisper")
    if not file.content_type.startswith("audio/"):
        raise HTTPException(400, "Только аудиофайлы")

    path = UPLOAD_DIR / f"{uuid.uuid4()}_{file.filename}"
    path.write_bytes(await file.read())
    try:
        model = get_whisper()
        segments, info = model.transcribe(str(path), language="ru")
        text = " ".join(seg.text.strip() for seg in segments)
        return {"text": text, "language": info.language, "duration": info.duration}
    finally:
        path.unlink(missing_ok=True)

@app.post("/analyze")
async def analyze(req: TextRequest):
    """Анализ текста и извлечение структурированных данных через LLM."""
    if req.doc_type == "protocol":
        system = SYSTEM_PROTOCOL
        prompt = PROMPT_PROTOCOL.format(text=req.text)
    else:
        system = SYSTEM_PASSPORT
        prompt = PROMPT_PASSPORT.format(text=req.text)

    raw = await llm_generate(prompt, system)

    # очистка JSON от возможных markdown-блоков
    raw = re.sub(r"```json\s*", "", raw)
    raw = re.sub(r"```\s*", "", raw)
    raw = raw.strip()

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        # fallback: вернуть сырой текст
        data = {"raw": raw, "parse_error": True}

    return {"data": data, "doc_type": req.doc_type}

@app.post("/chat")
async def chat(req: ChatRequest):
    """Диалоговый режим: уточнение данных у пользователя."""
    system = """Ты — дружелюбный корпоративный ИИ-помощник специалиста отдела улучшений.
Твоя задача — помочь пользователю собрать всю необходимую информацию для оформления документа.
Задавай уточняющие вопросы, если информации недостаточно. Общайся на русском языке, коротко и по делу."""

    messages = [{"role": m.role, "content": m.content} for m in req.messages]

    # вызов Ollama напрямую с историей
    if not OLLAMA_AVAILABLE:
        raise HTTPException(503, "httpx не установлен")
    async with httpx.AsyncClient(timeout=120) as client:
        r = await client.post(f"{OLLAMA_URL}/api/chat", json={
            "model": OLLAMA_MODEL,
            "messages": [{"role": "system", "content": system}] + messages,
            "stream": False,
            "options": {"temperature": 0.5, "num_predict": 500}
        })
        r.raise_for_status()
        reply = r.json()["message"]["content"]
    return {"reply": reply}

@app.post("/generate-docx")
async def generate_docx(
    data: str = Form(...),
    doc_type: str = Form("protocol")
):
    """Генерация DOCX-документа по структурированным данным."""
    if not DOCX_AVAILABLE:
        raise HTTPException(503, "python-docx не установлен: pip install python-docx")

    parsed = json.loads(data)
    doc = DocxDocument()

    # стили
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h.runs[0] if h.runs else h.add_run()
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        return h

    def add_para(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def add_table_2col(rows):
        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(rows):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = str(v)
            tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    now = datetime.now().strftime("%d.%m.%Y %H:%M")
    file_id = uuid.uuid4().hex[:8]

    if doc_type == "protocol":
        # ── ПРОТОКОЛ ────────────────────────────────────────────────────────
        add_heading("ПРОТОКОЛ СЕССИИ РЕШЕНИЯ ПРОБЛЕМ", level=1)
        doc.add_paragraph(f"Дата формирования: {now}")
        doc.add_paragraph()

        add_heading("1. Общие сведения", level=2)
        add_table_2col([
            ("Дата сессии", parsed.get("date", "—")),
            ("Участники", ", ".join(parsed.get("participants", [])) or "—"),
        ])
        doc.add_paragraph()

        add_heading("2. Описание проблемы", level=2)
        add_para(parsed.get("problem", "—"))
        doc.add_paragraph()

        add_heading("3. Корневые причины", level=2)
        for rc in parsed.get("root_causes", []):
            doc.add_paragraph(rc, style="List Bullet")
        doc.add_paragraph()

        add_heading("4. Предложенные решения", level=2)
        for sol in parsed.get("solutions", []):
            doc.add_paragraph(sol, style="List Bullet")
        doc.add_paragraph()

        add_heading("5. План действий", level=2)
        actions = parsed.get("actions", [])
        if actions:
            tbl = doc.add_table(rows=1+len(actions), cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Задача"
            hdr[1].text = "Ответственный"
            hdr[2].text = "Срок"
            for run in [c.paragraphs[0].runs[0] for c in hdr if c.paragraphs[0].runs]:
                run.bold = True
            for i, action in enumerate(actions):
                row = tbl.rows[i+1].cells
                row[0].text = action.get("task", "")
                row[1].text = action.get("responsible", "")
                row[2].text = action.get("deadline", "")
        doc.add_paragraph()

        add_heading("6. Ожидаемый эффект", level=2)
        add_para(parsed.get("expected_effect", "—"))
        doc.add_paragraph()

        add_heading("7. Следующие шаги", level=2)
        add_para(parsed.get("next_steps", "—"))

        filename = f"protocol_{file_id}.docx"

    else:
        # ── ПАСПОРТ ПРОЕКТА ──────────────────────────────────────────────────
        add_heading("ПАСПОРТ ПРОЕКТА", level=1)
        doc.add_paragraph(f"Дата формирования: {now}")
        doc.add_paragraph()

        add_heading("1. Общие сведения", level=2)
        add_table_2col([
            ("Краткое наименование", parsed.get("project_name", "—")),
            ("Полное наименование", parsed.get("full_name", "—")),
            ("Категория проекта", parsed.get("category", "—")),
            ("Руководитель проекта", parsed.get("project_manager", "—")),
            ("Базовый срок реализации", parsed.get("timeline", "—")),
        ])
        doc.add_paragraph()

        add_heading("2. Цель проекта", level=2)
        add_para(parsed.get("goal", "—"))
        doc.add_paragraph()

        add_heading("3. Описание проблемы", level=2)
        add_para(parsed.get("problem_description", "—"))
        doc.add_paragraph()

        add_heading("4. Текущее состояние (AS-IS)", level=2)
        add_para(parsed.get("current_state", "—"))
        doc.add_paragraph()

        add_heading("5. Целевое состояние (TO-BE)", level=2)
        add_para(parsed.get("target_state", "—"))
        doc.add_paragraph()

        add_heading("6. Экономический эффект", level=2)
        add_para(parsed.get("economic_effect", "—"))
        doc.add_paragraph()

        add_heading("7. Необходимые ресурсы", level=2)
        add_para(parsed.get("resources", "—"))
        doc.add_paragraph()

        add_heading("8. Риски", level=2)
        for risk in parsed.get("risks", []):
            doc.add_paragraph(risk, style="List Bullet")
        doc.add_paragraph()

        add_heading("9. Метрики успеха", level=2)
        for metric in parsed.get("success_metrics", []):
            doc.add_paragraph(metric, style="List Bullet")
        doc.add_paragraph()

        add_heading("10. Подписи", level=2)
        add_table_2col([
            ("Заказчик", ""),
            ("Куратор проекта", ""),
            ("Руководитель проекта", ""),
        ])

        filename = f"passport_{file_id}.docx"

    out_path = OUTPUT_DIR / filename
    doc.save(str(out_path))
    return FileResponse(str(out_path), filename=filename,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ── статика (фронтенд) ────────────────────────────────────────────────────────
frontend_path = Path("../frontend")
if frontend_path.exists():
    app.mount("/", StaticFiles(directory=str(frontend_path), html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
